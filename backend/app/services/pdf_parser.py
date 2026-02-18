import io
import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text_parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}, falling back to PyPDF2")
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e2:
            logger.error(f"PyPDF2 also failed: {e2}")
            return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        try:
            return file_bytes.decode("utf-8")
        except Exception:
            return ""


def detect_ats_issues(text: str) -> list[str]:
    issues = []
    lines = text.split("\n")
    short_lines = [l for l in lines if 1 < len(l.strip()) < 20]
    if len(short_lines) > len(lines) * 0.4:
        issues.append("Possible multi-column layout — ATS may misread this")
    special_bullets = re.findall(r"[●■◆▶◉★✓✗]", text)
    if len(special_bullets) > 5:
        issues.append("Special bullet characters found — replace with standard hyphens")
    word_count = len(text.split())
    if word_count < 200:
        issues.append("Resume appears very short — consider adding more detail")
    elif word_count > 1200:
        issues.append("Resume may be too long — aim for 1-2 pages")
    required_sections = ["experience", "education", "skills"]
    text_lower = text.lower()
    missing = [s for s in required_sections if s not in text_lower]
    if missing:
        issues.append(f"Missing standard sections: {', '.join(missing)}")
    return issues